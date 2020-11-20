# modifying docx

import docx
d = docx.Document('c:\\users\Mason\\documents\\demo.docx')
d.paragraphs

d.paragraphs[0].text
# 'Document Title'
d.paragraphs[1].text
# 'A plain paragraph having some bold and some italic.'
p = d.paragraphs[1]
p.runs
p.runs[0].text
# 'A plain paragraph having some '
p.runs[1].text
# 'bold'
p.runs[2].text
# ' and some '
p.runs[3].text
 #'italic.'
p.runs[1].bold
# True
p.runs[0].bold == None
# True
p.runs[3].italic
# True
p.runs[3].underline = True
p.runs[3].text = 'italic and underlined'

d.save('c:\\demo2.docx')
p.style
# 'Normal'
p.style = 'Title'
d.save('c:\\demo3.docx')

d = docx.Document() # to create a blank new document
d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('This is another paragraph.')
d.save('c:\\demo4.docx')
p = d.paragraphs[0]
p.add_run('This is a new run.')
p.runs
p.runs[1].bold = True
d.save('c:\\demo5.docx')
       
# To modify an existing word document, open it up and create a new document
docx.Document()

